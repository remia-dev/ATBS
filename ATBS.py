# Chapter 12 Web Scraping
# What is web scraping?

# Project mapIt.py
# Insert address in command line
    # import webbrowser, sys, pyperclip
    # if len(sys.argv) > 1:
    #     address = ' '.join(sys.argv[1:])
    # else:
    #     address = pyperclip.paste()

    # webbrowser.open('http:://www.google.com/maps/place' + address)

# Selenium is actually much easier to manage

#--------------------------------------------------------------------------------------------#

# Project OpenWeb.py (Self)
    # from selenium import webdriver
    #     # Browser.get('http://www.google.com/maps/place')

    #     # Open second tab
    #         # Browser.execute_script("window.open('about:blank','secondtab');")
    #         # Browser.switch_to.window("secondtab")
    #         # Browser.get('http://www.wikipedia.org')

    #     # Open third tab
    #         # Browser.execute_script("window.open('about:blank','thirdtab');")
    #         # Browser.switch_to.window("thirdtab")
    #         # Browser.get('http://www.codeforces.com')
    # media = input("Type of Media: ")
    # # Maybe create a dictionary and then loop so it is much easier rather than if else?
    # # Input selection is better then I'll just create multiple functions?
    # def mediaType(media):
    #     Browser = webdriver.Firefox()
    #     if media.lower() == "wiki":
    #         # First
    #         Browser.get('http://www.wikipedia.org')
    #         # Second
    #         Browser.execute_script("window.open('about:blank','secondtab');")
    #         Browser.switch_to.window("secondtab")
    #         Browser.get('https://wiki.archlinux.org/')
    #         # Third
    #         Browser.execute_script("window.open('about:blank','thirdtab');")
    #         Browser.switch_to.window("thirdtab")
    #         Browser.get("https://the-scp.foundation/")
    #     elif media.lower() == "socials":
    #         # First
    #         Browser.get('https://www.youtube.com/')
    #         # Second
    #         Browser.execute_script("window.open('about:blank','secondtab');")
    #         Browser.switch_to.window("secondtab")
    #         Browser.get('https://www.reddit.com/')
    #     else:
    #         print("None")

    # mediaType(media)

#--------------------------------------------------------------------------------------------#

# Project ExtractWebpage.py(Self)
    # import requests
    # from bs4 import BeautifulSoup

    # url = 'https://www.wikipedia.org/'
    # reqs = requests.get(url)
    # soup = BeautifulSoup(reqs.text, 'html.parser')

    # f = open("test.txt", "w")
    # for link in soup.find_all('a'):
    #     data = link.get('href')
    #     f.write(data)
    #     f.write("\n")

    # f.close()

#---------------------------------------------------------------------------------------------#


# Okay so to save a file? you need to write it

# What is beautifulSoup?
# To note
    # requests.get() to download the file
    # open() with 'wb' writes the file in binary
    # iter_content()
    # write()
    # close()

# Parsing HTML with bs4
    # We do this to read the text from html to text?
    # Okay this will be slightly need to be updated (Read the docs)
        # from selenium import webdriver
        # driver = webdriver.Firefox()
        # driver.get('http://www.inventwithpython.com')
        # linkElem = driver.find_element("link text", "Read Online for Free")
        # linkElem.click()

# Filling out and submitting forms
    # from selenium import webdriver
    # browser = webdriver.Firefox()
    # browser.get('https://login.metafilter.com')
    # userElem = browser.find_element("id", 'user_name')
    # userElem.send_keys('somerando@gmail.com')

    # passElem = browser.find_element("id", 'user_pass')
    # passElem.send_keys('ajshdsakjd213')
    # passElem.submit()

# Sending Special Keys
    # from selenium import webdriver
    # from selenium.webdriver.common.keys import Keys
    # browser = webdriver.Firefox()
    # browser.get('https://nostarch.com')
    # htmlElem = browser.find_element('tag name', 'html')
    # htmlElem.send_keys(Keys.END)
    # htmlElem.send_keys(Keys.HOME)
    # browser.refresh()
    # browser.implicitly_wait(10)
    # browser.quit()

#----------------------------------------------------------------------------------------------------------#

# Project Emailer.py
    # from selenium import webdriver
    # from selenium.webdriver.common.keys import Keys
    # # Get the email and password
    # username = input("Username: ")
    # password = input("Password: ")

    # # Setup selenium
    # browser = webdriver.Firefox()
    # browser.get('https://mail.google.com/mail')

    # # pass the email and password through the function
    # userId = browser.find_element('id', 'identifierId')
    # userId.send_keys(username)

    # # browser.click()
    # browser.find_element('id', 'identifierNext').click()


#----------------------------------------------------------------------------------------------------------#

# Project Image downloader
    # from selenium import webdriver
    # image = input("Type of image: ")
    # browser = webdriver.Firefox()
    # browser.get('https://imgur.com/')

    # category = browser.find_element('id', 'Searchbar')
    # category.send_keys(image)

#-----------------------------------------------------------------------------------------------------------#


# Project 2048
    # from selenium import webdriver
    # from selenium.webdriver.common.keys import Keys

    # browser = webdriver.Firefox()
    # browser.get('https://www.2048.org/')
    # Elem = browser.find_element('tag name', 'html')

    # while(True):
    #     Elem.send_keys(Keys.UP)
    #     Elem.send_keys(Keys.RIGHT)
    #     Elem.send_keys(Keys.DOWN)
    #     Elem.send_keys(Keys.LEFT)

#-----------------------------------------------------------------------------------------------------------#


# Project Link verification
    # import requests, sys, webbrowser, bs4
    # from selenium import webdriver
    # from bs4 import BeautifulSoup

    # url = input("Website: ")
    # reqs = requests.get(url)
    # soup = bs4.BeautifulSoup(reqs.text, 'html.parser')
    # linkElems = soup.select('.package-snippet')

    # for link in soup.find_all('a'):
    #     data = link.get('href')
    #     # If the first two strings are // then remove it otherwise open it
    #     if data[0] == '/':
    #         webbrowser.open(data[2:])
    #     else:
    #         webbrowser.open(data)

#---------------------------------------------------------------------------------------------------------#

# Chapter 13: Working with Excel Spreadsheets

# Reading Excel Documents
# Using pyxl module
    # import openpyxl
    # wb = openpyxl.load_workbook('example.xlsx')

# Getting sheets from workbook
    # sheet3 = wb['Sheet3']
    # print(sheet3)


# Getting cells from sheet
    # sheet = wb['Sheet1']
# Getting the cell
    # print(sheet['A1'])
# Value of the cell
    # c = sheet['B1']
    # print(sheet['A1'].value ,c.value, "\n")

    # print('Row %s, Column %s is %s' % (c.row, c.coordinate, c.value), "\n")
    # print(c.column, c.coordinate)

    # print(sheet.cell(row = 1, column = 2))
    # print(sheet.cell(row = 1, column = 2).value, "\n")
    # print(sheet.max_row, sheet.max_column, "\n")
    # for i in range(1, 8, 2):
    #     print(i , sheet.cell(row = i, column = 2).value)

# Converting Between Column letters and numbers
    # import openpyxl
    # from openpyxl.utils import get_column_letter, column_index_from_string
    # wb = openpyxl.load_workbook('example.xlsx')
    # print(get_column_letter(1), "\n")
    # print(column_index_from_string('B'))

# Getting rows and columns from the sheet
# import openpyxl
# wb = openpyxl.load_workbook('example.xlsx')
# sheet = wb['Sheet1']
# print(tuple(sheet['A1':'C3']))

# Number of rows
# for rowOfCellObject in sheet['A1': 'C7']:
    # Value and coordinate
    # for cellObj in rowOfCellObject:
    #     print(cellObj.coordinate, cellObj.value)
    # print('--- End of Row --- ')

# print(list(sheet.columns)[2])
# for cellObj in list(sheet.columns)[2]:
#     print(cellObj.value)

#-------------------------------------------------------------------------------------------------------------#

# Project Reading Data.py

    # import openpyxl
    # wb = openpyxl.load_workbook('files/censuspopdata.xlsx')
    # sheet = wb['Population by Census Tract']

    # countyData = {}
    # Read the data from excel spreadsheet
    # for row in range(2 , sheet.max_row + 1):
    #     state = sheet['B' + str(row)].value
    #     county = sheet['C' + str(row)].value
    #     pop = sheet['D' + str(row)].value
    #     print(state, county, pop, "\n")

    # Populating the data structure
    # Remember dictionaries? Each would have a key-value pair:
    # In this example it would be a dict(state)-key(county)-key(pop, tract)-value
    # So it would appear countyData[state][county]['tracts' or 'pop']

    # countyData.set_default(state, {})
    # countyData[state].set_default(county, {'tracts': 0, 'pop': 0})

    # countyData[state][county]['tracts'] += 1
    # countyData[state][county]['pop'] += int(pop)

    # Count the number of census tracts in each country


    # Counts the total population of each country


    # Prints the result
        # Open and read the cells of an excel document
        # Calculate the tract and pop. data and store it in a data structure
        # write the data structure to a text file with a .py extension using pprint module

#---------------------------------------------------------------------------------------------------------------#


# Chapter 14 Working with Google Sheets
# import ezsheets
# ss = ezsheets.Spreadsheet('ugtKzQAA34Opfjj5xbSPzqh9qH8Kv5n9u0TnQGwuNF4')



# Chapter 15 Working with PDFs and Word Documents

# PDF Documents

# Extracting Page from Pdfs
# import PyPDF2
# from PyPDF2 import PdfReader
    # File = PdfReader('files/meetingminutes.pdf', 'rb')
        # Assigning a page
            # page = File.pages[0]
        # Printing said page
            # print(page.extract_text())

# Decrypting PDFs
    # File = PdfReader('files/encrypted.pdf')

# Function to know if the pdf file is encrypted
    # print(File.is_encrypted)
# print(File.pages[0].extract_text())

    # print(File.decrypt('rosebud'))
    # print(File.pages[0].extract_text())

# Creating PDFs

# Mergings PDFs
    # from PyPDF2 import PdfReader
    # from PyPDF2 import PdfMerger

    # pdf1 = PdfReader('files/meetingminutes.pdf')
    # pdf2 = PdfReader('files/meetingminutes2.pdf')
    # pdfs = [pdf1, pdf2]

    # merger = PdfMerger()

    # for pdf in pdfs:
    #     merger.append(pdf)

    # merger.write("files/combinedmeetingminutes.pdf")

# Rotating Pages
    # from PyPDF2 import PdfReader, PdfWriter

    # File = PdfReader('files/meetingminutes.pdf')
    # writer = PdfWriter()

    # # Specify the file page cuz we are using a new file output
    # writer.add_page(File.pages[0])
    # # Rotate said file
    # writer.pages[0].rotate(180)


    # with open("output.pdf", "wb") as fp:
    #     writer.write(fp)

# Overlaying Pages

# Why can't i continuesly write using a for loop?
    # from PyPDF2 import PdfReader, PdfMerger, PdfWriter

    # pdf1 = PdfReader('files/meetingminutes.pdf', "rb")
    # overlay = PdfReader('files/watermark.pdf', "rb")
    # # page.mergePage(overlay)
    # writer = PdfWriter()

    # # I want it to write in a new file
    # for index in range(len(pdf1.pages)):
    #     p1 = pdf1.pages[index]
    #     p2 = overlay.pages[0]
    #     p1.merge_page(p2)
    #     writer.add_page(p1)


    # with open("watermarkedcover.pdf", "wb") as new:
    #     writer.write(new)

# Encrypting PDfs
# from pypdf import PdfReader, PdfWriter



#---------------------------------------------------------------------------------------------------------------------#


#----------------------------------------------------------------------------------------------------------------#

# Word Documents
import docx

# Reading Word Documents
doc = docx.Document('files/demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
